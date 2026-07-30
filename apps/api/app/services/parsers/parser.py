import csv
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass(slots=True)
class ParsedBlock:
    text: str
    page_number: int | None = None
    section_title: str | None = None
    paragraph_number: int | None = None
    sheet_name: str | None = None
    row_number: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(slots=True)
class ParsedDocument:
    blocks: list[ParsedBlock]
    page_count: int = 0
    source_type: str = "document"
    warnings: list[str] = field(default_factory=list)
    qa_items: list[dict[str, str]] = field(default_factory=list)


class UnsupportedDocumentError(ValueError):
    pass


class DocumentParser:
    supported = {".pdf", ".docx", ".txt", ".md", ".markdown", ".csv", ".xlsx", ".json"}

    def parse(self, path: Path) -> ParsedDocument:
        suffix = path.suffix.lower()
        if suffix not in self.supported:
            raise UnsupportedDocumentError(f"지원하지 않는 파일 형식입니다: {suffix}")
        return {
            ".pdf": self._pdf,
            ".docx": self._docx,
            ".txt": self._text,
            ".md": self._text,
            ".markdown": self._text,
            ".csv": self._csv,
            ".xlsx": self._xlsx,
            ".json": self._json,
        }[suffix](path)

    def _text(self, path: Path) -> ParsedDocument:
        last_error: UnicodeDecodeError | None = None
        for encoding in ("utf-8-sig", "utf-8", "cp949"):
            try:
                text = path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError as exc:
                last_error = exc
        else:
            raise ValueError("문서 인코딩을 확인할 수 없습니다.") from last_error
        blocks: list[ParsedBlock] = []
        section: str | None = None
        for index, paragraph in enumerate(filter(None, (p.strip() for p in text.split("\n\n"))), start=1):
            first = paragraph.splitlines()[0].strip()
            if first.startswith("#"):
                section = first.lstrip("#").strip()
            blocks.append(ParsedBlock(paragraph, section_title=section, paragraph_number=index))
        return ParsedDocument(blocks=blocks, page_count=1)

    def _pdf(self, path: Path) -> ParsedDocument:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        if reader.is_encrypted:
            raise ValueError("암호화된 PDF는 처리할 수 없습니다.")
        blocks: list[ParsedBlock] = []
        warnings: list[str] = []
        for page_index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                warnings.append(f"{page_index}페이지에 추출 가능한 텍스트가 없습니다. OCR이 필요할 수 있습니다.")
                continue
            for paragraph_index, paragraph in enumerate(filter(None, (p.strip() for p in text.split("\n\n"))), start=1):
                blocks.append(ParsedBlock(paragraph, page_number=page_index, paragraph_number=paragraph_index))
        if not blocks:
            raise ValueError("텍스트가 없는 스캔 PDF입니다. OCR 처리 후 다시 업로드해 주세요.")
        return ParsedDocument(blocks=blocks, page_count=len(reader.pages), warnings=warnings)

    def _docx(self, path: Path) -> ParsedDocument:
        from docx import Document

        document = Document(str(path))
        blocks: list[ParsedBlock] = []
        section: str | None = None
        paragraph_number = 0
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            paragraph_number += 1
            if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                section = text
            blocks.append(ParsedBlock(text, section_title=section, paragraph_number=paragraph_number))
        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(ParsedBlock(" | ".join(cells), section_title=f"표 {table_index}", row_number=row_index, metadata={"table_index": table_index}))
        return ParsedDocument(blocks=blocks, page_count=0)

    def _csv(self, path: Path) -> ParsedDocument:
        rows = self._read_csv(path)
        if rows and {"question", "answer"}.issubset({key.lower() for key in rows[0]}):
            normalized = [{key.lower(): value for key, value in row.items()} for row in rows]
            qa = [{"question": row.get("question", ""), "answer": row.get("answer", ""), "category": row.get("category", "일반")} for row in normalized if row.get("question") and row.get("answer")]
            blocks = [ParsedBlock(self.qa_embedding_text(item), row_number=index, section_title=item["category"], metadata={"source_type": "qa"}) for index, item in enumerate(qa, start=2)]
            return ParsedDocument(blocks=blocks, page_count=1, source_type="qa", qa_items=qa)
        blocks = [ParsedBlock(" | ".join(f"{key}: {value}" for key, value in row.items()), row_number=index) for index, row in enumerate(rows, start=2)]
        return ParsedDocument(blocks=blocks, page_count=1)

    def _read_csv(self, path: Path) -> list[dict[str, str]]:
        for encoding in ("utf-8-sig", "utf-8", "cp949"):
            try:
                with path.open("r", encoding=encoding, newline="") as handle:
                    return list(csv.DictReader(handle))
            except UnicodeDecodeError:
                continue
        raise ValueError("CSV 인코딩을 확인할 수 없습니다.")

    def _json(self, path: Path) -> ParsedDocument:
        data = json.loads(path.read_text(encoding="utf-8-sig"))
        rows = data if isinstance(data, list) else data.get("items", []) if isinstance(data, dict) else []
        if rows and all(isinstance(row, dict) and "question" in row and "answer" in row for row in rows):
            qa = [{"question": str(row["question"]), "answer": str(row["answer"]), "category": str(row.get("category", "일반"))} for row in rows]
            blocks = [ParsedBlock(self.qa_embedding_text(item), row_number=index, section_title=item["category"], metadata={"source_type": "qa"}) for index, item in enumerate(qa, start=1)]
            return ParsedDocument(blocks=blocks, page_count=1, source_type="qa", qa_items=qa)
        text = json.dumps(data, ensure_ascii=False, indent=2)
        return ParsedDocument(blocks=[ParsedBlock(text)], page_count=1)

    def _xlsx(self, path: Path) -> ParsedDocument:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        blocks: list[ParsedBlock] = []
        for sheet in workbook.worksheets:
            rows = sheet.iter_rows(values_only=True)
            header = [str(value or "") for value in next(rows, ())]
            for index, values in enumerate(rows, start=2):
                cells = [f"{header[col] or f'열 {col + 1}'}: {value}" for col, value in enumerate(values) if value is not None]
                if cells:
                    blocks.append(ParsedBlock(" | ".join(cells), sheet_name=sheet.title, row_number=index))
        return ParsedDocument(blocks=blocks, page_count=0)

    @staticmethod
    def qa_embedding_text(item: dict[str, str]) -> str:
        return f"질문: {item['question']}\n답변: {item['answer']}\n카테고리: {item.get('category', '일반')}"
